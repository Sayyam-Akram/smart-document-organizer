"""DOCX Text Extraction"""
from io import BytesIO
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        
        docx_file = BytesIO(docx_bytes)
        document = Document(docx_file)
        
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text_content.append(cell.text)
        
        full_text = "\n".join(text_content)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        return full_text
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""
